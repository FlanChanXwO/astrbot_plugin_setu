"""Docx 文件生成服务，用于封装 R18 图片。"""

from __future__ import annotations

import io
import os
from pathlib import Path

from astrbot.api import logger


class DocxService:
    """Docx 文件生成服务。"""

    def __init__(self):
        self._has_dependency = self._check_dependency()

    def _check_dependency(self) -> bool:
        """检查 python-docx 是否已安装。"""
        try:
            import importlib.util

            return importlib.util.find_spec("docx") is not None
        except ImportError:
            logger.error("python-docx 未安装，请运行: pip install python-docx")
            return False

    def create_docx_with_images(
        self,
        images: list[bytes],
        output_path: Path | None = None,
        tags: list[str] | None = None,
    ) -> Path | None:
        """创建包含图片的 Docx 文件。

        参数:
            images: 图片字节数据列表
            output_path: 输出文件路径，如果不指定则使用临时目录
            tags: 搜索标签列表，用于生成文件名

        返回:
            生成的 Docx 文件路径，失败返回 None
        """
        # 提前检查依赖，避免重复 try-import
        if not self._has_dependency:
            logger.error("python-docx 未安装，无法生成 Docx 文件")
            return None

        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Inches
        except ImportError:
            logger.error("python-docx 未安装，无法生成 Docx 文件")
            return None

        try:
            # 创建文档
            doc = Document()

            # 设置页面边距为窄边距，最大化图片显示区域
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)

            # 添加每张图片
            for i, img_data in enumerate(images):
                # 添加图片到文档
                img_stream = io.BytesIO(img_data)

                # 添加段落并插入图片
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                run = paragraph.add_run()
                # 设置图片宽度为页面宽度减去边距（约6英寸）
                run.add_picture(img_stream, width=Inches(6.0))

                # 如果不是最后一张图片，添加一些间距
                if i < len(images) - 1:
                    doc.add_paragraph()

            # 确定输出路径
            if output_path is None:
                from astrbot.core.utils.astrbot_path import get_astrbot_temp_path

                temp_dir = get_astrbot_temp_path()
                if isinstance(temp_dir, str):
                    temp_dir = Path(temp_dir)

                # 生成文件名：标签_随机UID.docx
                if tags:
                    # 清理标签，移除特殊字符
                    clean_tags = [
                        self._sanitize_filename(t) for t in tags[:3]
                    ]  # 最多取前3个标签
                    tag_str = ",".join(clean_tags)
                else:
                    tag_str = "setu"

                random_suffix = os.urandom(6).hex()  # 12位随机十六进制
                filename = f"{tag_str}_{random_suffix}.docx"
                output_path = temp_dir / filename

            # 保存文档
            doc.save(str(output_path))
            logger.info("Docx 文件已生成: %s", output_path)
            return output_path

        except (OSError, ValueError, RuntimeError) as e:
            logger.error("生成 Docx 文件失败: %s", e)
            return None

    def _sanitize_filename(self, text: str) -> str:
        """清理文件名中的非法字符。"""
        import re

        # 移除或替换Windows和Linux中的非法字符
        illegal_chars = r'[<>:"/\\|?*\x00-\x1f]'
        sanitized = re.sub(illegal_chars, "", text)

        # 限制长度，避免文件名过长
        if len(sanitized) > 30:
            sanitized = sanitized[:30]

        # 去除首尾空格和点
        sanitized = sanitized.strip(" .")

        # 如果为空，返回默认值
        if not sanitized:
            return "tag"

        return sanitized

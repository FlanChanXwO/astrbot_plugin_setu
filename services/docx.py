"""Docx 文件生成服务。"""

from __future__ import annotations

import io
import os
import re
from pathlib import Path

from astrbot.api import logger


class DocxService:
    """Docx 文件生成服务。"""

    def __init__(self):
        self._has_dependency = self._check_dependency()

    async def initialize(self) -> None:
        """初始化 Docx 服务（异步占位，实际初始化在 __init__ 完成）。"""
        pass

    def _check_dependency(self) -> bool:
        """检查 python-docx 是否已安装。"""
        try:
            import importlib.util

            return importlib.util.find_spec("docx") is not None
        except ImportError:
            logger.error("python-docx 未安装，请运行: pip install python-docx")
            return False

    def create_docx_with_images(
        self,
        images: list[bytes],
        output_path: Path | None = None,
        tags: list[str] | None = None,
    ) -> Path | None:
        """创建包含图片的 Docx 文件。"""
        if not self._has_dependency:
            logger.error("python-docx 未安装，无法生成 Docx 文件")
            return None

        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Inches
        except ImportError:
            logger.error("python-docx 未安装，无法生成 Docx 文件")
            return None

        try:
            doc = Document()

            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)

            for i, img_data in enumerate(images):
                img_stream = io.BytesIO(img_data)
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                run = paragraph.add_run()
                run.add_picture(img_stream, width=Inches(6.0))

                if i < len(images) - 1:
                    doc.add_paragraph()

            if output_path is None:
                from astrbot.core.utils.astrbot_path import get_astrbot_temp_path

                temp_dir = get_astrbot_temp_path()
                if isinstance(temp_dir, str):
                    temp_dir = Path(temp_dir)

                if tags:
                    clean_tags = [self._sanitize_filename(t) for t in tags[:3]]
                    tag_str = ",".join(clean_tags)
                else:
                    tag_str = "setu"

                random_suffix = os.urandom(6).hex()
                filename = f"{tag_str}_{random_suffix}.docx"
                output_path = temp_dir / filename

            doc.save(str(output_path))
            logger.info("Docx 文件已生成: %s", output_path)
            return output_path

        except (OSError, ValueError, RuntimeError) as e:
            logger.error("生成 Docx 文件失败: %s", e)
            return None

    def _sanitize_filename(self, text: str) -> str:
        """清理文件名中的非法字符。"""
        illegal_chars = r'[<>:"/\\|?*\x00-\x1f]'
        sanitized = re.sub(illegal_chars, "", text)

        if len(sanitized) > 30:
            sanitized = sanitized[:30]

        sanitized = sanitized.strip(" .")

        if not sanitized:
            return "tag"

        return sanitized
